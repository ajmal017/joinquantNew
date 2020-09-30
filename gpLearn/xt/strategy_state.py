# -*- coding: utf-8 -*-
'''
@Time    : 2020/9/25 17:23
@Author  : zhangfang
@File    : test.py
'''
import docx
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.shared import Inches
from docx.enum.table import WD_TABLE_ALIGNMENT
import copy
import datetime
import matplotlib.pyplot as plt
import pandas as pd
from backtest_func import yearsharpRatio, maxRetrace, annROR
from jqdatasdk import *
from configDB import *
auth(JOINQUANT_USER, JOINQUANT_PW)


def stock_price(sec, sday, eday):
    """
    输入 股票代码，开始日期，截至日期
    输出 个股的后复权的开高低收价格
    """
    temp = get_price(sec, start_date=sday, end_date=eday, frequency='daily', fields=['close'], skip_paused=True, fq=None,
                     count=None).close.tolist()[-1]
    return temp


def transfercode(x):
    x = str(x)
    if len(x) < 6:
        x = '0' * (6 - len(x)) + x
    if x[0] == '6':
        x = x + '.SH'
    else:
        x = x + '.SZ'
    return x


def trans_to_bs(x):
    if x[0] == '买':
        return 1
    else:
        return -1


def save_df_to_doc(document, test_df, word=None):
    """
    将结果按照dataframe的形式存入doc文件
    :param document: 存入的文档类
    :param test_df: 需要保存的df
    :return:

    """
    document.styles['Normal'].font.name = u'宋体'
    document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')

    # add_paragraph表示添加一个段落
    if word:
        document.add_paragraph(u'\n%s' % (word))
    if len(test_df.columns) == 0:
        # document.add_paragraph(u'\n%s' % ('无'))
        return

    # 添加一个表格--行数和列数，行数多加一行，需要将列名同时保存
    t = document.add_table(test_df.shape[0] + 1, test_df.shape[1], style="Table Grid")
    t.alignment = WD_TABLE_ALIGNMENT.CENTER  # 表格整体居中
    # 将每列列名保存到表格中
    for j in range(test_df.shape[-1]):
        t.cell(0, j).text = test_df.columns[j]
        t.cell(0, j).width = Inches(1.85)

    # 将每列数据保存到新建的表格中
    for i in range(test_df.shape[0]):
        for j in range(test_df.shape[-1]):
            # 第一行保存的是列名，所以数据保存时，行数要加1
            t.cell(i + 1, j).text = str(test_df.values[i, j])

    for row in t.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                for run in paragraph.runs:
                    font = run.font
                    font.size = Pt(7.5)
    for col in range(test_df.shape[1]):
        t.cell(0, col).width = Inches(1.65)


if __name__ == "__main__":
    fold_path = 'c:/e/data/qe/backtest/'
    first_date = '2020-09-18'
    end_date = first_date
    cash_ini = 500000
    strategy_name_lst = ['aiznqd', 'znlbtgj', 'sdxf', 'tlaefyh']
    strategy_china_name_lst = ['AI智能驱动', '智能罗伯特管家', '时代先锋', '淘利阿尔法1号']
    asset_df = pd.DataFrame(columns=['date'])
    state_lst = []
    state_value_lst = []
    for i in range(len(strategy_name_lst)):
        cash = 500000
        strategy_name = strategy_name_lst[i]
        china_name = strategy_china_name_lst[i]
        detail_zx = pd.read_excel(fold_path + strategy_name + '.xls', encoding='gbk')[
            ['代码', '操作时间', '交易费用', '市值', '业务类型', '操作价格', '数量']]
        detail_zx['date'] = detail_zx['操作时间'].apply(lambda x: str(x)[:10])
        detail_zx['time'] = detail_zx['操作时间'].apply(lambda x: str(x)[-8:])
        detail_zx['stock_code'] = detail_zx['代码'].apply(lambda x: transfercode(x))
        detail_zx['fee'] = detail_zx['交易费用']
        detail_zx['value'] = detail_zx['市值']
        detail_zx['bs'] = detail_zx['业务类型'].apply(lambda x: trans_to_bs(x))
        detail_zx['price'] = detail_zx['操作价格']
        detail_zx['volume'] = detail_zx['数量']
        detail_zx = detail_zx[['date', 'time', 'stock_code', 'fee', 'value', 'bs', 'price', 'volume']].sort_values(['date'])
        print(detail_zx)
        end_date = max(end_date, max(detail_zx.date))
        net = []
        net.append([first_date, cash_ini])
        for date, group in detail_zx.groupby(['date']):
            buy_df = group[group['bs'] == 1]
            sell_df = group[group['bs'] == -1]
            cash = cash + sell_df['value'].sum() - buy_df['value'].sum() - group['fee'].sum()
            stock_value = 0
            if len(buy_df) > 0:
                for code, group in buy_df.groupby(['stock_code']):
                    volume = group.volume.sum()
                    close = stock_price(normalize_code(code), date, date)
                    stock_value = stock_value + volume * close
            asset = cash + stock_value
            print(date, asset)
            net.append([date, asset])
        ret = pd.DataFrame(net, columns=['date', strategy_name])
        ret[china_name] = ret[strategy_name] / ret[strategy_name].tolist()[0]
        print(ret)
        asset_df = asset_df.merge(ret, on=['date'], how='outer')
        net_lst = ret[china_name].tolist()
        annR = annROR(net_lst, 1)
        sharp = yearsharpRatio(net_lst, 1)
        max_retrace = maxRetrace(net_lst, 1)
        print('sharp:%s' % sharp)
        print('annR:%s' % annR)
        print('max_retrace:%s' % max_retrace)
        state_value_lst.append(
            [first_date, end_date, china_name, cash_ini,
             ret[strategy_name].tolist()[-1], ret[strategy_name].tolist()[-1] - cash_ini])
        state_lst.append([first_date, end_date, china_name, net_lst[-1] - 1, annR, sharp, max_retrace])
    asset_df['date'] = pd.to_datetime(asset_df['date'])
    print(asset_df)

    asset_df_value = asset_df.set_index(['date']).ix[:, strategy_name_lst]
    asset_df_value['组合'] = asset_df_value.sum(axis=1)

    asset_df_value['组合净值'] = asset_df_value['组合'] / asset_df_value['组合'].tolist()[0]
    asset_df['组合'] = asset_df_value['组合净值'].tolist()
    print(asset_df)
    title_str = '策略净值曲线'
    name_lst = copy.deepcopy(strategy_china_name_lst)
    name_lst.append('组合')
    asset_df.set_index(['date']).ix[:, name_lst].plot()
    plt.rcParams['font.sans-serif'] = ['SimHei']
    plt.title(title_str)
    plt.savefig(fold_path + 'fig/' + 'net_' + end_date + '.png')

    net_lst = asset_df_value['组合净值'].tolist()
    annR = annROR(net_lst, 1)
    sharp = yearsharpRatio(net_lst, 1)
    max_retrace = maxRetrace(net_lst, 1)
    state_lst.append([first_date, end_date, '组合', net_lst[-1] - 1, annR, sharp, max_retrace])

    state_df = pd.DataFrame(state_lst, columns=['开始日期', '结束日期', '策略', '累计收益', '年化收益', '夏普', '最大回撤'])
    state_df['累计收益'] = state_df['累计收益'].apply(lambda x: '%.2f%%' % (x * 100))
    state_df['年化收益'] = state_df['年化收益'].apply(lambda x: '%.2f%%' % (x * 100))
    state_df['最大回撤'] = state_df['最大回撤'].apply(lambda x: '%.2f%%' % (x * 100))
    state_df['夏普'] = state_df['夏普'].apply(lambda x: '%.2f' % x)

    state_value_lst.append(
        [first_date, end_date, '组合', len(strategy_name_lst) * cash_ini,
         asset_df_value['组合'].tolist()[-1], asset_df_value['组合'].tolist()[-1] - len(strategy_name_lst) * cash_ini])
    state_value_df = pd.DataFrame(state_value_lst, columns=['开始日期', '结束日期', '策略', '初始资产', '当前资产', '盈亏'])
    state_value_df['初始资产'] = state_value_df['初始资产'].apply(lambda x: '%.2f' % x)
    state_value_df['当前资产'] = state_value_df['当前资产'].apply(lambda x: '%.2f' % x)
    state_value_df['盈亏'] = state_value_df['盈亏'].apply(lambda x: '%.2f' % x)

    document = Document()
    document.add_heading(f'模拟交易报告{end_date}', level=0)
    document.add_heading(f'策略汇总')
    save_df_to_doc(document, state_value_df, '账户资产')
    save_df_to_doc(document, state_df, '策略表现')

    document.add_heading(f'净值曲线')
    document.add_picture(f'{fold_path}/fig/net_{end_date}.png', width=Inches(6.0))

    document.save(f'{fold_path}/模拟交易报告{end_date}.docx')
