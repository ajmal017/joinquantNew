# -*- coding: utf-8 -*-
'''
@Time    : 2020/9/25 17:23
@Author  : zhangfang
@File    : test.py
'''
import docx
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.shared import Inches
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import RGBColor
from docx.enum.text import WD_LINE_SPACING, WD_PARAGRAPH_ALIGNMENT
import requests
import json
import copy
import datetime
import matplotlib.pyplot as plt
import pandas as pd
from backtest_func import yearsharpRatio, maxRetrace, annROR
from jqdatasdk import *
from configDB import *
import os
from logger_local import logger_self
plt.style.use('ggplot')
#table_style = 'Light List Accent 1'
table_style = "Table Grid"
auth(JOINQUANT_USER, JOINQUANT_PW)

fold_path = 'c:/e/data/qe/simulate/'
get_log = logger_self(set_level='info', file_path=fold_path + 'log' + '/').InItlogger()
strategy_map = {
    '淘利阿尔法1号': 'ptf8b5f6accedf11e9b8c20a580a81060a',
    '智能罗伯特管家': 'ptb99659f8cedf11e98c710a580a81060a',
    'AI智能驱动': 'pt608bd9c8cedf11e985790a580a81060a',
    '时代先锋': 'pt83dc374ccedf11e981d70a580a81060a',
}


def stock_price(sec, sday, eday):
    """
    输入 股票代码，开始日期，截至日期
    输出 个股的后复权的开高低收价格
    """
    temp = get_price(sec, start_date=sday, end_date=eday, frequency='daily', fields=['close'], skip_paused=True, fq=None,
                     count=None).close.tolist()[-1]
    return temp


def transfercode(x):
    x = str(x)
    if len(x) < 6:
        x = '0' * (6 - len(x)) + x
    if x[0] == '6':
        x = x + '.SH'
    else:
        x = x + '.SZ'
    return x


def trans_to_bs(x):
    if x[0] == '买':
        return 1
    else:
        return -1


def save_df_to_doc(document, test_df, word=None):
    """
    将结果按照dataframe的形式存入doc文件
    :param document: 存入的文档类
    :param test_df: 需要保存的df
    :return:

    """
    document.styles['Normal'].font.name = u'宋体'
    document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')

    # add_paragraph表示添加一个段落
    if word:
        document.add_paragraph(u'\n%s' % (word))
    if len(test_df.columns) == 0:
        # document.add_paragraph(u'\n%s' % ('无'))
        return

    # 添加一个表格--行数和列数，行数多加一行，需要将列名同时保存
    t = document.add_table(test_df.shape[0] + 1, test_df.shape[1], style=table_style)
    t.autofit = True
    t.alignment = WD_TABLE_ALIGNMENT.CENTER  # 表格整体居中
    # 将每列列名保存到表格中
    for j in range(test_df.shape[-1]):
        t.cell(0, j).text = test_df.columns[j]
        t.cell(0, j).width = Inches(1.85)

    # 将每列数据保存到新建的表格中
    for i in range(test_df.shape[0]):
        for j in range(test_df.shape[-1]):
            # 第一行保存的是列名，所以数据保存时，行数要加1
            t.cell(i + 1, j).text = str(test_df.values[i, j])

    for row in t.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                for run in paragraph.runs:
                    font = run.font
                    font.size = Pt(7.5)
    for col in range(test_df.shape[1]):
        t.cell(0, col).width = Inches(1.65)
        if t.cell(0, col).text == '策略名称':
            t.cell(0, col).width = Inches(3)


def get_date(calen, today):
    next_tradeday = get_trade_days(start_date=today + datetime.timedelta(days=1), end_date='2030-01-01')[0]
    # if datetime.datetime.now().hour >= 15:
    #     calen.append(next_tradeday)
    EndDate = calen[-1]
    StartDate = calen[0]
    hq_last_date = calen[-2]
    return calen, next_tradeday, EndDate, StartDate, str(hq_last_date)[:10]


def get_file_name(file_lst, date_name, field):
    for i in file_lst:
        if date_name in i and field in i:
            get_log.info('success get_file_name_%s: %s date:%s' % (field, i, date_name))
            return i
    get_log.info('get_file_name_account error: can not get %s_%s ' % (date_name, field))
    return False


def get_today_account(foldPath, file_lst, account_init, end_date):
    file_name_today_account = get_file_name(file_lst, end_date, 'account')
    if file_name_today_account:
        df_today = pd.read_csv(foldPath+file_name_today_account, encoding='gbk') \
                            .loc[:, ['交易日', '总资产', '总市值', '可用金额', '可取金额']]
        df_today['初始资产'] = account_init
        df_today['总资产'] = df_today['总资产'].apply(lambda x: '%.0f' % x)
        df_today['初始资产'] = df_today['初始资产'].apply(lambda x: '%.0f' % x)
        df_today['总市值'] = df_today['总市值'].apply(lambda x: '%.0f' % x)
        df_today['可用金额'] = df_today['可用金额'].apply(lambda x: '%.0f' % x)
        df_today['可取金额'] = df_today['可取金额'].apply(lambda x: '%.0f' % x)
        return df_today.loc[:, ['交易日', '初始资产', '总资产', '总市值', '可用金额', '可取金额']]
    else:
        return pd.DataFrame([], columns=['交易日', '初始资产', '总资产', '总市值', '可用金额', '可取金额'])


def get_strategy_account(foldPath, file_lst, end_date, strategy_name_lst, calen, init_asset):
    file_name_today = get_file_name(file_lst, end_date, 'all_today')
    all_today = pd.read_excel(foldPath+file_name_today, encoding='gbk') \
                    .loc[:, ['策略名称', '总买额', '总卖额', '净买额', '手续费', '盈亏']].dropna().set_index(['策略名称'])
    account = []
    asset_df = pd.DataFrame(calen, columns=['date'])
    for strategy_name in strategy_name_lst:
        file_name_history = get_file_name(file_lst, end_date, strategy_name + '_history')
        df_today = pd.read_excel(foldPath+file_name_history, encoding='gbk') \
                            .loc[:, ['策略名称', '日期', '当日成交额', '当日盈亏', '累计盈亏']]
        china_name = df_today['策略名称'].tolist()[0]
        df_today['日期'] = df_today['日期'].apply(lambda x: str(x)[:8])
        df_today = df_today[df_today['日期'] > calen[0]]
        asset_lst = [init_asset]
        asset_upgrate = init_asset
        for i in range(1, len(calen)-1):
            date = calen[i]
            df_ = df_today[df_today['日期'] == date]
            if len(df_) > 0:
                asset_upgrate = asset_upgrate + df_['当日盈亏'].tolist()[0]
            asset_lst.append(asset_upgrate)
        today_asset = all_today.loc[china_name]['盈亏'] + asset_upgrate
        asset_lst.append(today_asset)
        net_lst = [i/asset_lst[0] for i in asset_lst]
        annR = annROR(net_lst, 1)
        sharp = yearsharpRatio(net_lst, 1)
        max_retrace = maxRetrace(net_lst, 1)
        total_ret = net_lst[-1] - 1
        today_ret = net_lst[-1] / net_lst[-2] - 1
        account.append([china_name, init_asset, today_asset, total_ret, today_ret, annR, sharp, max_retrace])
        asset_df[china_name] = net_lst
        asset_df[strategy_name] = asset_lst
    asset_df['date'] = pd.to_datetime(asset_df['date'])
    asset_df_value = asset_df.set_index(['date']).ix[:, strategy_name_lst]
    asset_df_value['组合'] = asset_df_value.sum(axis=1)

    asset_df_value['组合净值'] = asset_df_value['组合'] / asset_df_value['组合'].tolist()[0]
    asset_df['组合'] = asset_df_value['组合净值'].tolist()

    title_str = '策略净值曲线'
    name_lst = copy.deepcopy(strategy_china_name_lst)
    name_lst.append('组合')
    asset_df.set_index(['date']).ix[:, name_lst].plot()
    plt.rcParams['font.sans-serif'] = ['SimHei']
    plt.title(title_str)
    plt.savefig(fold_path + 'fig/' + 'net_' + end_date + '.png')

    net_lst = asset_df_value['组合净值'].tolist()
    today_asset = asset_df_value['组合'].tolist()[-1]
    total_ret = net_lst[-1] - 1
    today_ret = net_lst[-1] / net_lst[-2] - 1
    annR = annROR(net_lst, 1)
    sharp = yearsharpRatio(net_lst, 1)
    max_retrace = maxRetrace(net_lst, 1)
    account.append(['组合', asset_df_value['组合'].tolist()[0], today_asset, total_ret, today_ret, annR, sharp, max_retrace])

    account_df = pd.DataFrame(account, columns=['策略名称', '初始资产', '当前资产', '总收益', '当日收益', '年化收益', '夏普', '最大回撤'])
    all_today['总买额'] = all_today['总买额'].apply(lambda x: '%.0f' % x)
    all_today['总卖额'] = all_today['总卖额'].apply(lambda x: '%.0f' % x)
    all_today['净买额'] = all_today['净买额'].apply(lambda x: '%.0f' % x)
    all_today['手续费'] = all_today['手续费'].apply(lambda x: '%.2f' % x)
    all_today['当日盈亏'] = all_today['盈亏'].apply(lambda x: '%.0f' % x)
    all_today['策略名称'] = all_today.index
    # account_df['开始日期'] = calen[0]
    # account_df['结束日期'] = calen[-1]
    account_df['当前资产'] = account_df['当前资产'].apply(lambda x: '%.0f' % x)
    account_df['夏普'] = account_df['夏普'].apply(lambda x: '%.2f' % x)
    account_df['年化收益'] = account_df['年化收益'].apply(lambda x: '%.2f%%' % (x * 100))
    account_df['总收益'] = account_df['总收益'].apply(lambda x: '%.2f%%' % (x * 100))
    account_df['当日收益'] = account_df['当日收益'].apply(lambda x: '%.2f%%' % (x * 100))
    account_df['最大回撤'] = account_df['最大回撤'].apply(lambda x: '%.2f%%' % (x * 100))
    print(account_df)

    return all_today[['策略名称', '总买额', '总卖额', '净买额', '手续费', '当日盈亏']], account_df[
        ['策略名称', '初始资产', '当前资产', '总收益', '当日收益', '年化收益', '夏普', '最大回撤']]


def get_today_hold(foldPath, file_lst, end_date):
    file_name_today_hold = get_file_name(file_lst, end_date, 'hold')
    if file_name_today_hold:
        df_today = pd.read_csv(foldPath+file_name_today_hold, encoding='gbk') \
                            .loc[:, ['证券代码', '证券名称', '当前拥股', '持仓成本', '成本价', '市值', '盈亏比例']]
        df_today['证券代码'] = df_today['证券代码'].apply(lambda x: transfercode(x))
        df_today['持股'] = df_today['当前拥股'].apply(lambda x: '%.0f' % x)
        df_today['持仓成本'] = df_today['持仓成本'].apply(lambda x: '%.0f' % x)
        df_today['成本价'] = df_today['成本价'].apply(lambda x: '%.2f' % x)
        df_today['市值'] = df_today['市值'].apply(lambda x: '%.0f' % x)
        df_today['盈亏比例'] = df_today['盈亏比例'].apply(lambda x: '%.2f%%' % (x))
        return df_today.loc[:, ['证券代码', '证券名称', '持股', '持仓成本', '成本价', '市值', '盈亏比例']]
    else:
        return pd.DataFrame([], columns=['证券代码', '证券名称', '持股', '持仓成本', '成本价', '市值', '盈亏比例'])


def get_today_trade(fold_path, file_lst, end_date):
    file_name_today_hold = get_file_name(file_lst, end_date, 'trade_detail')
    if file_name_today_hold:
        df_today = pd.read_excel(fold_path+file_name_today_hold, encoding='gbk') \
                            .loc[:, ['证券代码', '证券名称', '操作', '成交价格', '成交数量', '成交金额', '手续费', '成交次数']]
        df_today['证券代码'] = df_today['证券代码'].apply(lambda x: transfercode(x))
        df_today['成交数量'] = df_today['成交数量'].apply(lambda x: '%.0f' % x)
        df_today['手续费'] = df_today['手续费'].apply(lambda x: '%.0f' % x)
        df_today['成交次数'] = df_today['成交次数'].apply(lambda x: '%.0f' % x)
        df_today['成交价格'] = df_today['成交价格'].apply(lambda x: '%.2f' % x)
        df_today['成交金额'] = df_today['成交金额'].apply(lambda x: '%.0f' % x)
        return df_today.loc[:, ['证券代码', '证券名称', '操作', '成交价格', '成交数量', '成交金额', '手续费', '成交次数']]
    else:
        return pd.DataFrame([], columns=['证券代码', '证券名称', '操作', '成交价格', '成交数量', '成交金额', '手续费', '成交次数'])


def get_account_state(fold_path, file_lst, calen):
    net_df = []
    for date in calen:
        for i in file_lst:
            if date in i and 'account' in i:
                df_ = pd.read_csv(fold_path + i, encoding='gbk') \
                       .loc[:, ['交易日', '总资产']]
                net_df.append(df_)
                break
    net_df = pd.concat(net_df).sort_values(['交易日'])
    print(net_df)
    state_sdate = calen[0]
    state_edate = calen[-1]
    asset_lst = net_df['总资产'].tolist()
    net_lst = [i/asset_lst[0] for i in asset_lst]
    annR = annROR(net_lst, 1)
    sharp = yearsharpRatio(net_lst, 1)
    max_retrace = maxRetrace(net_lst, 1)
    total_ret = net_lst[-1] - 1
    today_ret = net_lst[-1] / net_lst[-2] - 1
    print('sharp:%s' % sharp)
    print('annR:%s' % annR)
    print('max_retrace:%s' % max_retrace)
    state_value_lst = []
    state_value_lst.append([
        state_sdate, state_edate, asset_lst[0], asset_lst[-1], total_ret, today_ret, annR, sharp, max_retrace])
    df_today = pd.DataFrame(
        state_value_lst, columns=[
            '开始日期', '结束日期', '初始资产', '当前资产', '总收益', '当日收益', '年化收益', '夏普', '最大回撤'])
    df_today['初始资产'] = df_today['初始资产'].apply(lambda x: '%.0f' % x)
    df_today['当前资产'] = df_today['当前资产'].apply(lambda x: '%.0f' % x)
    df_today['夏普'] = df_today['夏普'].apply(lambda x: '%.2f' % x)
    df_today['年化收益'] = df_today['年化收益'].apply(lambda x: '%.2f%%' % (x * 100))
    df_today['总收益'] = df_today['总收益'].apply(lambda x: '%.2f%%' % (x * 100))
    df_today['当日收益'] = df_today['当日收益'].apply(lambda x: '%.2f%%' % (x * 100))
    df_today['最大回撤'] = df_today['最大回撤'].apply(lambda x: '%.2f%%' % (x * 100))
    return df_today


def planedorders(s_date, e_date, unique_id):
    url = 'https://aicloud.citics.com/papertrading/api/obtain_orders'
    data = {
        # 可选，不填时读取所有策略
        'unique_id': unique_id,  # 'pt461a4de0cedf11e995950a580a81060a' ,
        # 必填，<API文档>中的API_KEY
        'api_key': 'eyJ0eXAiOiJKV1QiLCJhbGciOiJIUzI1NiJ9.eyJ1c2VyX2lkIjoxLCJleHAiOjcyNTgwODk2MDAsInVwZGF0ZV90aW1lIjoiMjAyMC0wOC0yOSAxNDowMDo1NCIsInVzZXJuYW1lIjoiYnFhZG0xIn0.mHwdmeHTW7vchJjKt1b_iiKwbJ8L-Soq6wy0g607ri4',
        'start_date': s_date[:10],
        'end_date': e_date[:10],
        'username': 'bqadm1',
    }
    r = requests.post(url=url, data=data)
    res_dict = json.loads(r.text)
    print(res_dict)
    if res_dict['code'] != 200:
        print('获取信号失败，code： %s' % res_dict['code'])
        return [], []
    else:
        if len(res_dict['data']['order_list']) == 0:
            print('当日无信号')
            return [], []
        else:
            df = pd.DataFrame.from_dict(res_dict['data']['order_list'])[['run_date', 'symbol', 'trade_side']]
            df_buy = df[df['trade_side'] == 1]
            df_sell = df[df['trade_side'] == 2]
            if len(df_buy) == 0:
                buy_symbol = []
            else:
                buy_symbol = df_buy.symbol.tolist()
            if len(df_sell) == 0:
                sell_symbol = []
            else:
                sell_symbol = df_sell.symbol.tolist()
            return buy_symbol, sell_symbol


def get_strategy_signal(strategy_name_lst, date):
    lst = []
    for strategy_name in strategy_name_lst:
        buy_markets, sell_markets = planedorders(date, date, strategy_map[strategy_name])
        lst.append([strategy_name, date, buy_markets, sell_markets])
    ret = pd.DataFrame(lst, columns=['策略名称', '日期', '当日买入信号', '当日卖出信号'])
    return ret


if __name__ == "__main__":
    strategy_name_lst = ['aiznqd', 'znlbtgj', 'sdxf', 'tlaefyh']
    strategy_china_name_lst = ['AI智能驱动', '智能罗伯特管家', '时代先锋', '淘利阿尔法1号']
    # strategy_name_lst = ['aiznqd', 'znlbtgj']
    # strategy_china_name_lst = ['AI智能驱动', '智能罗伯特管家']
    account_init = 10000000
    strategy_init = 500000
    s_date = '2020-09-30'
    today = datetime.date.today()
    calen = get_trade_days(s_date, today)
    calen = [i.strftime('%Y%m%d') for i in list(calen)]
    print(calen)
    calen, next_tradeday, EndDate, StartDate, hq_last_date = get_date(calen, today)
    get_log.info('EndDate:%s, StartDate:%s' % (EndDate, StartDate))
    document = Document()
    # 添加标题,并修改字体样式
    head = document.add_heading(0)
    run = head.add_run(f'模拟盘交易报告-{EndDate}')
    run.font.name = u'黑体'  # 设置字体为黑体
    run._element.rPr.rFonts.set(qn('w:eastAsia'), u'黑体')
    run.font.size = Pt(24)  # 设置大小为24磅
    run.font.color.rgb = RGBColor(0, 0, 0)  # 设置颜色为黑色
    head.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # 居中

    fold_path_account = fold_path + 'account/'
    fold_path_hold = fold_path + 'hold/'
    fold_path_trade = fold_path + 'trade/'
    fold_path_strategy = fold_path + 'strategy/'

    file_lst_account = os.listdir(fold_path_account)
    file_lst_hold = os.listdir(fold_path_hold)
    file_lst_trade = os.listdir(fold_path_trade)
    file_lst_strategy = os.listdir(fold_path_strategy)

    get_log.info('写入总账户信息')
    document.add_heading(f'总账户信息')
    # 账户总览
    today_account = get_today_account(fold_path_account, file_lst_account, account_init, EndDate)
    save_df_to_doc(document, today_account, '账户总览')
    # 账户统计
    account_state = get_account_state(fold_path_account, file_lst_account, calen)
    save_df_to_doc(document, account_state, '账户统计')
    # 账户持仓
    today_hold = get_today_hold(fold_path_hold, file_lst_hold, EndDate)
    save_df_to_doc(document, today_hold, '当前持仓')
    # 当日成交
    today_trade = get_today_trade(fold_path_trade, file_lst_trade, EndDate)
    save_df_to_doc(document, today_trade, '当日交易')
    get_log.info('写入总账户信息完成')

    get_log.info('写入策略信息')
    document.add_heading(f'策略信息')
    strategy_today_trade, strategy_state = get_strategy_account(fold_path_strategy,
        file_lst_strategy, EndDate, strategy_name_lst, calen, strategy_init)
    save_df_to_doc(document, strategy_state, '策略统计%s-%s' %(StartDate, EndDate))
    document.add_heading(f'净值曲线')
    document.add_picture(f'{fold_path}/fig/net_{EndDate}.png', width=Inches(6.0))
    signal_today = get_strategy_signal(strategy_china_name_lst, EndDate[:4] + '-' + EndDate[4:6] + '-' + EndDate[6:])
    save_df_to_doc(document, signal_today, '当日交易信号')
    save_df_to_doc(document, strategy_today_trade, '当日成交信息')
    document.save(f'{fold_path}/report/模拟交易报告{EndDate}.docx')
