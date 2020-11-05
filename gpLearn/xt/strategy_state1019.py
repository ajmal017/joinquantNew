# -*- coding: utf-8 -*-
'''
@Time    : 2020/9/25 17:23
@Author  : zhangfang
@File    : test.py
'''
import docx
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.shared import Inches
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import RGBColor
from docx.enum.text import WD_LINE_SPACING, WD_PARAGRAPH_ALIGNMENT
import requests
import json
import copy
import datetime
import matplotlib.pyplot as plt
import pandas as pd
from backtest_func import yearsharpRatio, maxRetrace, annROR
from jqdatasdk import *
from configDB import *
import os
from logger_local import logger_self
from gpLearn.xt.word_pdf import FormatConvert, SendMessage
import numpy as np
plt.style.use('ggplot')
#table_style = 'Light List Accent 1'
table_style = "Table Grid"
auth(JOINQUANT_USER, JOINQUANT_PW)

# fold_path = 'c:/e/simulate/'
fold_path = 'c:/e/data/qe/simulate/'
get_log = logger_self(set_level='info', file_path=fold_path + 'log' + '/').InItlogger()
strategy_map = {
    '淘利阿尔法1号': 'ptf8b5f6accedf11e9b8c20a580a81060a',
    '智能罗伯特管家': 'ptb99659f8cedf11e98c710a580a81060a',
    'AI智能驱动': 'pt608bd9c8cedf11e985790a580a81060a',
    '时代先锋': 'pt83dc374ccedf11e981d70a580a81060a',
}
def stock_price_index(sec, period, sday, eday):
    """
    输入 股票代码，开始日期，截至日期
    输出 个股的后复权的开高低收价格
    """
    temp = get_price(sec, start_date=sday, end_date=eday, frequency=period,
                     skip_paused=False, fq='pre', count=None).reset_index() \
        .rename(columns={'index': 'trade_date'})\
        .assign(trade_date=lambda df: df.trade_date.apply(lambda x: str(x)[:10])).dropna()
    temp['stock_code'] = sec
    return temp

def stock_price(sec, sday, eday):
    """
    输入 股票代码，开始日期，截至日期
    输出 个股的后复权的开高低收价格
    """
    temp = get_price(sec, start_date=sday, end_date=eday, frequency='daily', fields=['close'], skip_paused=True, fq=None,
                     count=None).close.tolist()[-1]
    return temp


def transfercode(x):
    x = str(x)
    if len(x) < 6:
        x = '0' * (6 - len(x)) + x
    if x[0] == '6':
        x = x[:6] + '.SH'
    else:
        x = x[:6] + '.SZ'
    return x


def trans_to_bs(x):
    if x[0] == '买':
        return 1
    else:
        return -1


def save_df_to_doc(document, test_df, word=None):
    """
    将结果按照dataframe的形式存入doc文件
    :param document: 存入的文档类
    :param test_df: 需要保存的df
    :return:

    """
    document.styles['Normal'].font.name = u'宋体'
    document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')

    # add_paragraph表示添加一个段落
    if word:
        document.add_paragraph(u'\n%s' % (word))
    if len(test_df.columns) == 0:
        # document.add_paragraph(u'\n%s' % ('无'))
        return

    # 添加一个表格--行数和列数，行数多加一行，需要将列名同时保存
    t = document.add_table(test_df.shape[0] + 1, test_df.shape[1], style=table_style)
    t.autofit = True
    t.alignment = WD_TABLE_ALIGNMENT.CENTER  # 表格整体居中
    # 将每列列名保存到表格中
    for j in range(test_df.shape[-1]):
        t.cell(0, j).text = test_df.columns[j]
        t.cell(0, j).width = Inches(1.85)

    # 将每列数据保存到新建的表格中
    for i in range(test_df.shape[0]):
        for j in range(test_df.shape[-1]):
            # 第一行保存的是列名，所以数据保存时，行数要加1
            t.cell(i + 1, j).text = str(test_df.values[i, j])

    for row in t.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                for run in paragraph.runs:
                    font = run.font
                    font.size = Pt(7.5)
    for col in range(test_df.shape[1]):
        t.cell(0, col).width = Inches(1.65)
        if t.cell(0, col).text == '策略名称':
            t.cell(0, col).width = Inches(3)


def get_date(calen, today):
    next_tradeday = get_trade_days(start_date=today + datetime.timedelta(days=1), end_date='2030-01-01')[0]
    # if datetime.datetime.now().hour >= 15:
    #     calen.append(next_tradeday)
    EndDate = calen[-1]
    StartDate = calen[0]
    hq_last_date = calen[-2]
    return calen, next_tradeday, EndDate, StartDate, str(hq_last_date)[:10]


def get_file_name(file_lst, date_name, field):
    for i in file_lst:
        if date_name in i and field in i:
            get_log.info('success get_file_name_%s: %s date:%s' % (field, i, date_name))
            return i
    get_log.info('get_file_name_account error: can not get %s_%s ' % (date_name, field))
    return False


def get_today_account(foldPath, file_lst, account_init, end_date):
    file_name_today_account = get_file_name(file_lst, end_date, 'account')
    if file_name_today_account:
        df_today = pd.read_csv(foldPath+file_name_today_account, encoding='gbk') \
                            .loc[:, ['交易日', '总资产', '总市值', '可用金额', '可取金额']]
        df_today['初始资产'] = account_init
        df_today['总资产'] = df_today['总资产'].apply(lambda x: '%.0f' % x)
        df_today['初始资产'] = df_today['初始资产'].apply(lambda x: '%.0f' % x)
        df_today['总市值'] = df_today['总市值'].apply(lambda x: '%.0f' % x)
        df_today['可用金额'] = df_today['可用金额'].apply(lambda x: '%.0f' % x)
        df_today['可取金额'] = df_today['可取金额'].apply(lambda x: '%.0f' % x)
        return df_today.loc[:, ['交易日', '初始资产', '总资产', '总市值', '可用金额', '可取金额']]
    else:
        return pd.DataFrame([], columns=['交易日', '初始资产', '总资产', '总市值', '可用金额', '可取金额'])


def get_strategy_trade(foldPath, end_date):
    signal_today_df = pd.read_csv(foldPath + 'signal_' + str(end_date) + '.csv', encoding='gbk')\
                          .loc[:, ['策略名称', '证券代码', '数量', '操作', '时间']]
    signal_today_df['证券代码'] = signal_today_df['证券代码'].apply(lambda x: str(x)[:9])
    signal_today_df = signal_today_df.groupby(['策略名称', '证券代码', '操作']).sum()


def get_strategy_tradedetail(foldPath, end_date):
    signal_today = pd.read_csv()



def get_strategy_account(foldPath, file_lst, end_date, strategy_name_lst, calen, init_asset):
    file_name_today = get_file_name(file_lst, end_date, 'all_today')
    all_today = pd.read_excel(foldPath+file_name_today, encoding='gbk') \
                    .loc[:, ['策略名称', '总买额', '总卖额', '净买额', '手续费', '盈亏']].dropna().set_index(['策略名称'])
    account = []
    asset_df = pd.DataFrame(calen, columns=['date'])
    for strategy_name in strategy_name_lst:
        file_name_history = get_file_name(file_lst, end_date, strategy_name + '_history')
        df_today = pd.read_excel(foldPath+file_name_history, encoding='gbk') \
                            .loc[:, ['策略名称', '日期', '当日成交额', '当日盈亏', '累计盈亏']]
        china_name = df_today['策略名称'].tolist()[0]
        df_today['日期'] = df_today['日期'].apply(lambda x: str(x)[:8])
        df_today = df_today[df_today['日期'] > calen[0]]
        asset_lst = [init_asset]
        asset_upgrate = init_asset
        for i in range(1, len(calen)-1):
            date = calen[i]
            df_ = df_today[df_today['日期'] == date]
            if len(df_) > 0:
                asset_upgrate = asset_upgrate + df_['当日盈亏'].tolist()[0]
            asset_lst.append(asset_upgrate)
        today_asset = all_today.loc[china_name]['盈亏'] + asset_upgrate
        asset_lst.append(today_asset)
        net_lst = [i/asset_lst[0] for i in asset_lst]
        annR = annROR(net_lst, 1)
        sharp = yearsharpRatio(net_lst, 1)
        max_retrace = maxRetrace(net_lst, 1)
        total_ret = net_lst[-1] - 1
        today_ret = net_lst[-1] / net_lst[-2] - 1
        account.append([china_name, init_asset, today_asset, total_ret, today_ret, annR, sharp, max_retrace])
        asset_df[china_name] = net_lst
        asset_df[strategy_name] = asset_lst
    asset_df['date'] = pd.to_datetime(asset_df['date'])
    asset_df_value = asset_df.set_index(['date']).ix[:, strategy_name_lst]
    asset_df_value['组合'] = asset_df_value.sum(axis=1)

    asset_df_value['组合净值'] = asset_df_value['组合'] / asset_df_value['组合'].tolist()[0]
    asset_df['组合'] = asset_df_value['组合净值'].tolist()

    title_str = '策略净值曲线'
    name_lst = copy.deepcopy(strategy_china_name_lst)
    name_lst.append('组合')
    asset_df.set_index(['date']).ix[:, name_lst].plot()
    plt.rcParams['font.sans-serif'] = ['SimHei']
    plt.title(title_str)
    plt.savefig(fold_path + 'fig/' + 'net_' + end_date + '.png')

    net_lst = asset_df_value['组合净值'].tolist()
    today_asset = asset_df_value['组合'].tolist()[-1]
    total_ret = net_lst[-1] - 1
    today_ret = net_lst[-1] / net_lst[-2] - 1
    annR = annROR(net_lst, 1)
    sharp = yearsharpRatio(net_lst, 1)
    max_retrace = maxRetrace(net_lst, 1)
    account.append(['组合', asset_df_value['组合'].tolist()[0], today_asset, total_ret, today_ret, annR, sharp, max_retrace])

    account_df = pd.DataFrame(account, columns=['策略名称', '初始资产', '当前资产', '总收益', '当日收益', '年化收益', '夏普', '最大回撤'])
    all_today['总买额'] = all_today['总买额'].apply(lambda x: '%.0f' % x)
    all_today['总卖额'] = all_today['总卖额'].apply(lambda x: '%.0f' % x)
    all_today['净买额'] = all_today['净买额'].apply(lambda x: '%.0f' % x)
    all_today['手续费'] = all_today['手续费'].apply(lambda x: '%.2f' % x)
    all_today['当日盈亏'] = all_today['盈亏'].apply(lambda x: '%.0f' % x)
    all_today['策略名称'] = all_today.index
    # account_df['开始日期'] = calen[0]
    # account_df['结束日期'] = calen[-1]
    account_df['当前资产'] = account_df['当前资产'].apply(lambda x: '%.0f' % x)
    account_df['夏普'] = account_df['夏普'].apply(lambda x: '%.2f' % x)
    account_df['年化收益'] = account_df['年化收益'].apply(lambda x: '%.2f%%' % (x * 100))
    account_df['总收益'] = account_df['总收益'].apply(lambda x: '%.2f%%' % (x * 100))
    account_df['当日收益'] = account_df['当日收益'].apply(lambda x: '%.2f%%' % (x * 100))
    account_df['最大回撤'] = account_df['最大回撤'].apply(lambda x: '%.2f%%' % (x * 100))
    print(account_df)

    return all_today[['策略名称', '总买额', '总卖额', '净买额', '手续费', '当日盈亏']], account_df[
        ['策略名称', '初始资产', '当前资产', '总收益', '当日收益', '年化收益', '夏普', '最大回撤']]


def get_date_hold(foldPath, file_lst, end_date):
    '''
    :param foldPath:
    :param file_lst:
    :param end_date:
    :return: 获取某一交易日持仓
    '''
    file_name_today_hold = get_file_name(file_lst, end_date, 'hold')
    if file_name_today_hold:
        df_today = pd.read_csv(foldPath+file_name_today_hold, encoding='gbk') \
                            .loc[:, ['证券代码', '证券名称', '当前拥股', '持仓成本', '成本价', '市值', '盈亏比例', '最新价']]\
            .rename(columns={'当前拥股': '持股'})
        print(df_today)
        if len(df_today) == 0:
            return df_today
        df_today['证券代码'] = df_today['证券代码'].apply(lambda x: transfercode(x))
        df_today[['持股', '持仓成本', '成本价', '市值', '盈亏比例', '最新价']] = df_today[['持股', '持仓成本', '成本价', '市值', '盈亏比例', '最新价']]\
            .astype(float)

        return df_today.loc[:, ['证券代码', '证券名称', '持股', '持仓成本', '成本价', '市值', '盈亏比例', '最新价']]
    else:
        return pd.DataFrame([], columns=['证券代码', '证券名称', '持股', '持仓成本', '成本价', '市值', '盈亏比例', '最新价'])


def get_trade_detail(fold_path, file_lst, end_date):
    file_name_today_hold = get_file_name(file_lst, end_date, 'trade')
    if file_name_today_hold:
        df_today = pd.read_csv(fold_path+file_name_today_hold, encoding='gbk') \
                            .loc[:, ['任务来源', '证券代码', '证券名称', '操作', '成交价格', '成交数量', '成交金额', '手续费',
                                     '成交日期', '成交时间']] \
            .rename(columns={'任务来源': '策略名称', '成交日期': '日期', '成交时间': '时间'})
        df_today['证券代码'] = df_today['证券代码'].apply(lambda x: transfercode(x))
        df_today[['成交价格', '成交数量', '成交金额', '手续费']] = df_today[['成交价格', '成交数量', '成交金额', '手续费']] \
            .astype(float)

        return df_today.loc[:, ['策略名称', '证券代码', '证券名称', '操作', '成交价格', '成交数量', '成交金额', '手续费',
                                     '日期', '时间']]
    else:
        return pd.DataFrame([], columns=['策略名称', '证券代码', '证券名称', '操作', '成交价格', '成交数量', '成交金额', '手续费',
                                     '日期', '时间'])


def get_trade_state(trade_detail):
    if len(trade_detail) == 0:
        return pd.DataFrame([], columns=['策略名称', '证券代码', '证券名称', '操作', '成交均价', '成交数量', '成交金额', '手续费', '日期'])
    trade_detail = trade_detail.loc[:, ['策略名称', '证券代码', '证券名称', '操作', '成交数量', '成交金额', '手续费', '日期']]
    df_today = trade_detail.groupby(['策略名称', '证券代码', '证券名称', '操作', '日期']).sum()\
        .reset_index(drop=False).sort_values(['日期', '策略名称', '操作'])
    df_today['成交均价'] = df_today['成交金额'] / df_today['成交数量']
    df_today['日期'] = df_today['日期'].apply(lambda x: str(x))
    return df_today[['策略名称', '证券代码', '证券名称', '操作', '成交均价', '成交数量', '成交金额', '手续费', '日期']]


def get_account_state(fold_path, file_lst, calen):
    net_df = []
    for date in calen:
        for i in file_lst:
            if date in i and 'account' in i:
                df_ = pd.read_csv(fold_path + i, encoding='gbk') \
                       .loc[:, ['交易日', '总资产']]
                net_df.append(df_)
                break
    net_df = pd.concat(net_df).sort_values(['交易日'])
    print(net_df)
    state_sdate = calen[0]
    state_edate = calen[-1]
    asset_lst = net_df['总资产'].tolist()
    net_lst = [i/asset_lst[0] for i in asset_lst]
    annR = annROR(net_lst, 1)
    sharp = yearsharpRatio(net_lst, 1)
    max_retrace = maxRetrace(net_lst, 1)
    total_ret = net_lst[-1] - 1
    today_ret = net_lst[-1] / net_lst[-2] - 1
    print('sharp:%s' % sharp)
    print('annR:%s' % annR)
    print('max_retrace:%s' % max_retrace)
    state_value_lst = []
    state_value_lst.append([
        state_sdate, state_edate, asset_lst[0], asset_lst[-1], total_ret, today_ret, annR, sharp, max_retrace])
    df_today = pd.DataFrame(
        state_value_lst, columns=[
            '开始日期', '结束日期', '初始资产', '当前资产', '总收益', '当日收益', '年化收益', '夏普', '最大回撤'])
    df_today['初始资产'] = df_today['初始资产'].apply(lambda x: '%.0f' % x)
    df_today['当前资产'] = df_today['当前资产'].apply(lambda x: '%.0f' % x)
    df_today['夏普'] = df_today['夏普'].apply(lambda x: '%.2f' % x)
    df_today['年化收益'] = df_today['年化收益'].apply(lambda x: '%.2f%%' % (x * 100))
    df_today['总收益'] = df_today['总收益'].apply(lambda x: '%.2f%%' % (x * 100))
    df_today['当日收益'] = df_today['当日收益'].apply(lambda x: '%.2f%%' % (x * 100))
    df_today['最大回撤'] = df_today['最大回撤'].apply(lambda x: '%.2f%%' % (x * 100))
    return df_today


def planedorders(s_date, e_date, unique_id):
    url = 'https://aicloud.citics.com/papertrading/api/obtain_orders'
    data = {
        # 可选，不填时读取所有策略
        'unique_id': unique_id,  # 'pt461a4de0cedf11e995950a580a81060a' ,
        # 必填，<API文档>中的API_KEY
        'api_key': 'eyJ0eXAiOiJKV1QiLCJhbGciOiJIUzI1NiJ9.eyJ1c2VyX2lkIjoxLCJleHAiOjcyNTgwODk2MDAsInVwZGF0ZV90aW1lIjoiMjAyMC0wOC0yOSAxNDowMDo1NCIsInVzZXJuYW1lIjoiYnFhZG0xIn0.mHwdmeHTW7vchJjKt1b_iiKwbJ8L-Soq6wy0g607ri4',
        'start_date': s_date[:10],
        'end_date': e_date[:10],
        'username': 'bqadm1',
    }
    r = requests.post(url=url, data=data)
    res_dict = json.loads(r.text)
    print(res_dict)
    if res_dict['code'] != 200:
        print('获取信号失败，code： %s' % res_dict['code'])
        return [], []
    else:
        if len(res_dict['data']['order_list']) == 0:
            print('当日无信号')
            return [], [], pd.DataFrame([])
        else:
            df = pd.DataFrame.from_dict(res_dict['data']['order_list'])[['run_date', 'symbol', 'trade_side']]
            df_buy = df[df['trade_side'] == 1]
            df_sell = df[df['trade_side'] == 2]
            if len(df_buy) == 0:
                buy_symbol = []
            else:
                buy_symbol = df_buy.symbol.tolist()
            if len(df_sell) == 0:
                sell_symbol = []
            else:
                sell_symbol = df_sell.symbol.tolist()
            return buy_symbol, sell_symbol, df


def get_strategy_signal_history(strategy_name_lst, s_date, e_date):
    lst = []
    for strategy_name in strategy_name_lst:
        buy_markets, sell_markets, df = planedorders(s_date, e_date, strategy_map[strategy_name])
        df['策略名称'] = strategy_name
        lst.append(df)
    ret = pd.concat(lst)
    return ret


def get_sell_plan(foldPath, file_lst, last_date):
    '''
    卖出计划：获取上一交易日持仓
    :param foldPath:
    :param file_lst:
    :param last_date:
    :return:
    '''
    last_hold_df = get_date_hold(foldPath, file_lst, last_date)[
        ['证券代码', '证券名称', '持股']]
    if len(last_hold_df) > 0:
        last_hold_df['持股'] = last_hold_df['持股'].apply(lambda x: int(x))
        last_hold_df = last_hold_df[last_hold_df['持股'] > 0]
        #last_hold_df['证券代码'] = last_hold_df['证券代码'].apply(lambda x: transfercode(x))
    return last_hold_df


def calculate_strategy_net(calen, next_tradeday, trade_fold_path, trade_file_lst, strategy_fold_path, account_hold, strategy_name_lst):
    trade_detail_all = []
    for date in calen:
        trade_detail = get_trade_detail(trade_fold_path, trade_file_lst, date)
        trade_detail_all.append(trade_detail)
    trade_detail_all = pd.concat(trade_detail_all)
    trade_state_all = get_trade_state(trade_detail_all)[
            ['策略名称', '证券代码', '证券名称', '操作', '成交均价', '成交数量', '成交金额', '手续费', '日期']]
    trade_state_all['bs'] = trade_state_all['操作'].apply(lambda x: trans_to_bs(x))
    trade_state_all['持仓'] = trade_state_all['成交数量'] * trade_state_all['bs']
    trade_state_all['成交金额'] = -trade_state_all['成交金额'] * trade_state_all['bs']
    win_r_lst = []

    for strategy_name in strategy_name_lst:
        temp = trade_state_all[trade_state_all['策略名称'] == strategy_name]
        print(temp)
        ret_lst = []
        if len(temp) == 0:
            win_r_lst.append([strategy_name, 1, 1])
            continue
        win_time = 0
        trad_time = 0
        cost_lst = []
        buy_code_lst = []
        for date, group in temp.groupby(['日期']):
            if len(buy_code_lst) > 0:
                for i in range(len(buy_code_lst)):
                    cost_i = cost_lst[i]
                    code = buy_code_lst[i]
                    sell_df = group[(group['bs'] == -1) & (group['证券代码'] == code)]
                    if len(sell_df) == 0:
                        continue
                    sell_value = abs(sell_df['成交金额'].sum()) - abs(sell_df['手续费'].sum())
                    ret = sell_value - cost_i
                    ret_lst.append(ret)
                    if ret > 0:
                        win_time += 1
                        trad_time += 1
                    else:
                        trad_time += 1
            cost_lst = []
            buy_code_lst = []
            buy_df = group[group['bs'] == 1]
            if len(buy_df) == 0:
                buy_code_lst = []
                continue
            for buy_code, buy_group in buy_df.groupby(['证券代码']):
                buy_code_lst.append(buy_code)
                cost = abs(buy_group['成交金额'].sum()) + abs(buy_group['手续费'].sum())
                cost_lst.append(cost)
        win_r = 1
        odd = 1
        if trad_time > 0:
            win_r = win_time / trad_time
            ret_pos = [i for i in ret_lst if i > 0]
            ret_nev = [i for i in ret_lst if i < 0]
            if len(ret_nev) > 0:
                pos_ave = 0
                if len(ret_pos) > 0:
                    pos_ave = abs(np.mean(ret_pos))
                odd = pos_ave / abs(np.mean(ret_nev))
        win_r_lst.append([strategy_name, win_r, odd])
    win_df = pd.DataFrame(win_r_lst, columns=['策略名称', '胜率', '盈亏比'])


    trade_state_today = trade_state_all[trade_state_all['日期'] == calen[-1]]
    hold_df_firstday = trade_state_all[(trade_state_all['日期'] == calen[0]) & (trade_state_all['bs'] == 1)]
    hold_df_back = trade_state_all[(trade_state_all['日期'] > calen[0])]
    hold_df = pd.concat([hold_df_firstday, hold_df_back])[['策略名称', '证券代码', '证券名称', '持仓']]\
        .groupby(['策略名称', '证券代码', '证券名称']).sum().reset_index(drop=False)

    print('==============hold_df')
    print(hold_df)
    hold_today = hold_df[hold_df['持仓'] > 0]
    hold_today = trade_state_all[['策略名称', '证券代码', '证券名称', '日期', '成交均价', 'bs']]\
        .merge(hold_today, on=['策略名称', '证券代码', '证券名称'])
    hold_today = hold_today[(hold_today['日期'] == hold_today['日期'].max()) & (hold_today['bs'] == 1)]

    account_hold_ = account_hold[['证券代码', '持股', '最新价']]
    account_hold_lst = account_hold['证券代码'].tolist()
    account_hold_ = account_hold_.set_index(['证券代码'])
    print('==============hold_today')
    print(hold_today)
    print('==============account_hold_')
    print(account_hold_)
    hold_ret_lst = []
    for strategy_name, group in hold_today.groupby(['策略名称']):
        for id, row_ in group.iterrows():
            code = row_['证券代码']
            if code not in account_hold_lst:
                continue
            hold = row_['持仓']
            price = account_hold_.loc[code]['最新价']
            cost__ = row_['成交均价']
            hold = min(hold, account_hold_.loc[code]['持股'])
            hold_ret_lst.append([strategy_name, code, row_['证券名称'], hold, cost__, price])
    hold_strategy_today = pd.DataFrame(hold_ret_lst, columns=['策略名称', '证券代码', '证券名称', '持仓', '成本价', '最新价'])
    print('==============hold_strategy_today')
    print(hold_strategy_today)

    strategy_account_all_today = []
    strategy_account_all_dict = {}

    for strategy_name in strategy_name_lst:
        name = strategy_dict[strategy_name]
        hold = hold_strategy_today[hold_strategy_today['策略名称'] == strategy_name]
        hold.index = hold['证券代码']
        trade_state = trade_state_today[trade_state_today['策略名称'] == strategy_name]
        last_strategy_account = pd.read_csv(
            strategy_fold_path + name + '_' + calen[-1] + '.csv', encoding='gbk')[
            ['策略名称', 'date', 'cash', 'position_value']]
        last_strategy_account['date'] = last_strategy_account['date'].apply(lambda x: str(x))
        last_strategy_account = last_strategy_account.set_index(['date'])
        #print(type(last_strategy_account['date'].tolist()[0]))

        last_cash = last_strategy_account.loc[calen[-2]]['cash']
        today_cash = last_cash + trade_state['成交金额'].sum() - trade_state['手续费'].sum()
        today_hold = 0
        print(hold)
        if len(hold) > 0:
            for code in hold['证券代码'].tolist():
                price = hold.loc[code]['最新价']
                position = hold.loc[code]['持仓']
                today_hold = today_hold + price * position
        today_strategy_account = pd.DataFrame([[strategy_name, str(calen[-1]), today_cash, today_hold]],
                                              columns=['策略名称', 'date', 'cash', 'position_value'])
        print(today_strategy_account)
        last_strategy_account = last_strategy_account.reset_index(drop=False)[['策略名称', 'date', 'cash', 'position_value']]
        print(last_strategy_account)
        strategy_account_new = pd.concat([last_strategy_account, today_strategy_account])
        print(strategy_account_new)
        next_tradeday = str(next_tradeday).replace('-', '')
        strategy_account_new.to_csv(strategy_fold_path + name + '_' + next_tradeday + '.csv', encoding='gbk')
        strategy_account_all_today.append(today_strategy_account)
        strategy_account_all_dict[strategy_name] = strategy_account_new
    strategy_account_all_today = pd.concat(strategy_account_all_today).sort_values(['策略名称', 'date'])
    strategy_state = []
    strategy_net_all_dict = {}
    profolio_asset_dict = {}
    for strategy_name in strategy_account_all_dict:
        df = strategy_account_all_dict[strategy_name]
        df['asset'] = df['cash'] + df['position_value']
        df['net'] = df['asset'] / df['asset'].tolist()[0]
        net_lst = df['net'].tolist()
        strategy_net_all_dict[strategy_name] = net_lst
        profolio_asset_dict[strategy_name] = df['asset'].tolist()
        today_profit = df['asset'].tolist()[-1] - df['asset'].tolist()[-2]
        today_ret = net_lst[-1] / net_lst[-2] - 1
        today_net = net_lst[-1]
        total_ret = today_net - 1
        annR = annROR(net_lst, 1)
        sharp = yearsharpRatio(net_lst, 1)
        max_retrace = maxRetrace(net_lst, 1)
        strategy_state.append([strategy_name, calen[0], calen[-1], today_net, total_ret, annR, sharp, max_retrace, today_profit, today_ret])
    profolio_asset_dict['date'] = calen
    profolio_asset_df = pd.DataFrame(profolio_asset_dict).set_index(['date'])
    profolio_asset_df['组合资产'] = profolio_asset_df.sum(axis=1)
    profolio_asset_df['组合'] = profolio_asset_df['组合资产'] / profolio_asset_df['组合资产'].tolist()[0]
    net_lst = profolio_asset_df['组合'].tolist()
    strategy_net_all_dict['组合'] = net_lst
    today_profit = profolio_asset_df['组合资产'].tolist()[-1] - profolio_asset_df['组合资产'].tolist()[-2]
    today_ret = net_lst[-1] / net_lst[-2] - 1
    today_net = net_lst[-1]
    total_ret = today_net - 1
    annR = annROR(net_lst, 1)
    sharp = yearsharpRatio(net_lst, 1)
    max_retrace = maxRetrace(net_lst, 1)
    strategy_state.append(
        ['组合', calen[0], calen[-1], today_net, total_ret, annR, sharp, max_retrace, today_profit, today_ret])
    strategy_state_df = pd.DataFrame(strategy_state, columns=['策略名称', '开始日期', '结束日期', '净值', '累计收益', '年化收益', '夏普', '最大回撤', '当日盈亏', '当日收益'])
    strategy_state_df = strategy_state_df.merge(win_df, on=['策略名称'])

    asset_df = pd.DataFrame(strategy_net_all_dict)
    asset_df['date'] = calen
    asset_df['date'] = pd.to_datetime(asset_df['date'])
    asset_df['沪深300'] = index_hq_net_lst

    title_str = '策略净值曲线'
    asset_df.set_index(['date']).plot()
    plt.rcParams['font.sans-serif'] = ['SimHei']
    plt.title(title_str)
    plt.savefig(fold_path + 'fig/' + 'net_' + calen[-1] + '.png')

    return strategy_account_all_today, hold_strategy_today, strategy_state_df


def get_strategy_signal(strategy_name_lst, date):
    lst = []
    for strategy_name in strategy_name_lst:
        buy_markets, sell_markets, df = planedorders(date, date, strategy_map[strategy_name])

        lst.append([strategy_name, date, buy_markets, sell_markets])
    ret = pd.DataFrame(lst, columns=['策略名称', '日期', '当日买入信号', '当日卖出信号'])
    return ret


if __name__ == "__main__":
    strategy_name_lst = ['aiznqd', 'znlbtgj', 'sdxf', 'tlaefyh']
    strategy_china_name_lst = ['AI智能驱动', '智能罗伯特管家', '时代先锋', '淘利阿尔法1号']
    strategy_dict = {'AI智能驱动': 'aiznqd', '智能罗伯特管家': 'znlbtgj', '时代先锋': 'sdxf', '淘利阿尔法1号': 'tlaefyh'}
    index_code = '000300.XSHG'

    # strategy_name_lst = ['aiznqd', 'znlbtgj']
    # strategy_china_name_lst = ['AI智能驱动', '智能罗伯特管家']
    account_init = 10208522
    strategy_init = 500000
    s_date = '2020-10-14'

    today = datetime.date.today()
    today = pd.to_datetime('2020-10-29')

    calen = get_trade_days(s_date, today)
    calen = [i.strftime('%Y%m%d') for i in list(calen)]
    print(calen)
    calen, next_tradeday, EndDate, StartDate, hq_last_date = get_date(calen, today)
    get_log.info('EndDate:%s, StartDate:%s' % (EndDate, StartDate))
    index_hq = stock_price_index(index_code, '1d', pd.to_datetime(StartDate), pd.to_datetime(EndDate))
    index_hq_net_lst = [i/index_hq['close'].tolist()[0] for i in index_hq['close'].tolist()]
    document = Document()
    # 添加标题,并修改字体样式
    head = document.add_heading(0)
    run = head.add_run(f'模拟盘交易报告-{EndDate}')
    run.font.name = u'黑体'  # 设置字体为黑体
    run._element.rPr.rFonts.set(qn('w:eastAsia'), u'黑体')
    run.font.size = Pt(24)  # 设置大小为24磅
    run.font.color.rgb = RGBColor(0, 0, 0)  # 设置颜色为黑色
    head.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # 居中

    fold_path_account = fold_path + 'account/'
    fold_path_hold = fold_path + 'hold/'
    fold_path_trade = fold_path + 'trade/'
    fold_path_strategy = fold_path + 'strategy/'
    fold_path_signal = fold_path + 'signal/'

    file_lst_account = os.listdir(fold_path_account)
    file_lst_hold = os.listdir(fold_path_hold)
    file_lst_trade = os.listdir(fold_path_trade)
    file_lst_strategy = os.listdir(fold_path_strategy)

    get_log.info('写入总账户信息')
    document.add_heading(f'总账户信息')
    # 账户总览
    today_account = get_today_account(fold_path_account, file_lst_account, account_init, EndDate)
    save_df_to_doc(document, today_account, '账户总览')
    # 账户统计
    account_state = get_account_state(fold_path_account, file_lst_account, calen)
    save_df_to_doc(document, account_state, '账户统计')
    # 账户持仓
    today_hold = get_date_hold(fold_path_hold, file_lst_hold, EndDate)
    strategy_account_all_today, hold_strategy_today, strategy_state_df = calculate_strategy_net(
        calen, next_tradeday, fold_path_trade, file_lst_trade, fold_path_strategy, today_hold, strategy_china_name_lst)
    if len(today_hold) > 0:
        today_hold = today_hold[today_hold['持股'] > 0]
        today_hold['持股'] = today_hold['持股'].apply(lambda x: '%.0f' % x)
        today_hold['持仓成本'] = today_hold['持仓成本'].apply(lambda x: '%.0f' % x)
        today_hold['成本价'] = today_hold['成本价'].apply(lambda x: '%.2f' % x)
        today_hold['市值'] = today_hold['市值'].apply(lambda x: '%.0f' % x)
        today_hold['盈亏比例'] = today_hold['盈亏比例'].apply(lambda x: '%.2f%%' % (x))
    save_df_to_doc(document, today_hold, '当前持仓')
    get_log.info('写入总账户信息完成')

    get_log.info('写入策略信息')
    document.add_heading(f'策略信息')
    strategy_account_all_today = strategy_account_all_today.rename(
        columns={'cash': '可用资金', 'position_value': '持仓市值'})
    if len(strategy_account_all_today) > 0:
        strategy_account_all_today['净资产'] = strategy_account_all_today['持仓市值'] + strategy_account_all_today['可用资金']
        strategy_account_all_today['持仓市值'] = strategy_account_all_today['持仓市值'].apply(lambda x: '%.0f' % x)
        strategy_account_all_today['可用资金'] = strategy_account_all_today['可用资金'].apply(lambda x: '%.0f' % x)
        strategy_account_all_today['净资产'] = strategy_account_all_today['净资产'].apply(lambda x: '%.0f' % x)
        strategy_account_all_today['初始资产'] = strategy_init

    save_df_to_doc(document, strategy_account_all_today[['策略名称', '持仓市值', '可用资金', '净资产', '初始资产']], '策略账户概览')

    if len(hold_strategy_today) > 0:
        hold_strategy_today['成本价'] = hold_strategy_today['成本价'].apply(lambda x: '%.2f' % x)
        hold_strategy_today['持仓'] = hold_strategy_today['持仓'].apply(lambda x: '%.0f' % x)
    save_df_to_doc(document, hold_strategy_today, '策略持仓')
    next_tradeday = str(next_tradeday).replace('-', '')
    hold_strategy_today.to_csv(fold_path_hold + 'hold_all_' + next_tradeday + '.csv', encoding='gbk')
    if len(hold_strategy_today) > 0:
        strategy_state_df['当日盈亏'] = strategy_state_df['当日盈亏'].apply(lambda x: '%.0f' % x)
        # strategy_state_df['净值'] = strategy_state_df['净值'].apply(lambda x: '%.2f' % x)
        strategy_state_df['夏普'] = strategy_state_df['夏普'].apply(lambda x: '%.2f' % x)
        strategy_state_df['年化收益'] = strategy_state_df['年化收益'].apply(lambda x: '%.2f%%' % (x*100))
        strategy_state_df['累计收益'] = strategy_state_df['累计收益'].apply(lambda x: '%.2f%%' % (x*100))
        strategy_state_df['最大回撤'] = strategy_state_df['最大回撤'].apply(lambda x: '%.2f%%' % (x*100))
        strategy_state_df['盈亏比'] = strategy_state_df['盈亏比'].apply(lambda x: '%.2f' % x)
        strategy_state_df['胜率'] = strategy_state_df['胜率'].apply(lambda x: '%.2f%%' % (x * 100))
        strategy_state_df['当日收益'] = strategy_state_df['当日收益'].apply(lambda x: '%.2f%%' % (x*100))
    save_df_to_doc(document, strategy_state_df.drop(['净值'], axis=1), '策略统计')
    # strategy_today_trade, strategy_state = get_strategy_account(fold_path_strategy,
    #     file_lst_strategy, EndDate, strategy_name_lst, calen, strategy_init)
    # save_df_to_doc(document, strategy_state, '策略统计%s-%s' %(StartDate, EndDate))
    document.add_heading(f'净值曲线')
    document.add_picture(f'{fold_path}/fig/net_{EndDate}.png', width=Inches(6.0))
    signal_today = get_strategy_signal(strategy_china_name_lst, EndDate[:4] + '-' + EndDate[4:6] + '-' + EndDate[6:])
    save_df_to_doc(document, signal_today, '当日买卖信号')
    sell_plan_df = get_sell_plan(fold_path_hold, file_lst_hold, hq_last_date)
    save_df_to_doc(document, sell_plan_df, '计划卖出')

    # 当日成交
    trade_detail = get_trade_detail(fold_path_trade, file_lst_trade, EndDate)
    today_trade_state = get_trade_state(trade_detail)

    if len(trade_detail) > 0:
        trade_detail['成交数量'] = trade_detail['成交数量'].apply(lambda x: '%.0f' % x)
        trade_detail['手续费'] = trade_detail['手续费'].apply(lambda x: '%.0f' % x)
        trade_detail['成交价格'] = trade_detail['成交价格'].apply(lambda x: '%.2f' % x)
        trade_detail['成交金额'] = trade_detail['成交金额'].apply(lambda x: '%.0f' % x)

        today_trade_state['成交数量'] = today_trade_state['成交数量'].apply(lambda x: '%.0f' % x)
        today_trade_state['手续费'] = today_trade_state['手续费'].apply(lambda x: '%.0f' % x)
        today_trade_state['成交均价'] = today_trade_state['成交均价'].apply(lambda x: '%.2f' % x)
        today_trade_state['成交金额'] = today_trade_state['成交金额'].apply(lambda x: '%.0f' % x)

    save_df_to_doc(document, today_trade_state[
        ['策略名称', '证券代码', '证券名称', '操作', '成交均价', '成交数量', '成交金额', '手续费']], '当日成交信息')

    # save_df_to_doc(document, trade_detail, '当日成交明细')
    document.save(f'{fold_path}/report/模拟盘交易报告{EndDate}.docx')

    FormatConvert.word_to_pdf(f'{fold_path}/report/模拟盘交易报告{EndDate}.docx',
                              f'{fold_path}/report/模拟盘交易报告{EndDate}.pdf')
    subject = f'模拟盘交易报告{EndDate}'

    password = '9eFzgacCkDMUpPP6'
    sender = 'aiquant@ai-quants.com'
    # 收件人为多个收件人
    # receiver = ["dawn0zhou@163.com", "zhangfang@ai-quants.com", "wjm@ai-quants.com",
    #             "zj@ai-quants.com", 'hzn@ai-quants.com']
    receiver = ['aiquant@ai-quants.com']
    send = SendMessage(sender, password)
    file_path = f'{fold_path}/report/模拟盘交易报告{EndDate}.pdf'
    # send.send_email(subject, subject, receiver, file_path)
